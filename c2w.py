import subprocess
import sys
import os
import time
import re
import tempfile
import platform
import shlex
from pathlib import Path # Using pathlib for robust path handling
from typing import List, Dict, Tuple, Any

# Core document libraries - import at startup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Remove webdriver imports from global scope - they'll be imported when needed
# This helps reduce startup time

# --- Configuration for Frontend Screenshots ---
BROWSER_BINARY_LOCATION = None
VIEWPORT_HEIGHT = 1394


# --- HELPER FUNCTIONS FOR INTERACTIVE EXECUTION ---

def is_interactive(code_content: str, lang_ext: str) -> bool:
    """
    Checks if the code content for a given language is likely to require user input.
    """
    if lang_ext == '.py':
        return bool(re.search(r'\binput\s*\(', code_content))
    elif lang_ext in ['.c', '.cpp']:
        # Look for scanf, gets, or cin usage
        return bool(re.search(r'\b(scanf|gets)\s*\(|\b(std::)?cin\b', code_content))
    elif lang_ext == '.java':
        return bool(re.search(r'new\s+Scanner\s*\(\s*System\.in\s*\)', code_content))
    return False


def open_interactive_terminal(command_list: List[str], working_dir: str, log_file_path: str) -> bool:
    """
    Opens a new OS-native terminal to run an interactive command, logs the entire
    session to log_file_path, and waits for it to complete.
    """
    system = platform.system()
    command_str = " ".join(shlex.quote(c) for c in command_list)
    print(f"Opening interactive session for: {command_str}")

    try:
        if system == "Windows":
            ps_command = (
                f"Start-Transcript -Path '{log_file_path}' -Force; "
                f"cd '{working_dir}'; "
                f"{command_str}; "
                f"pause; "
                f"Stop-Transcript"
            )
            subprocess.run(
                ['start', 'Interactive Session', '/wait', 'powershell', '-ExecutionPolicy', 'Bypass', '-Command', ps_command],
                shell=True, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
        elif system in ["Linux", "Darwin"]:
            # The 'script' command logs terminal sessions on Linux/macOS
            script_inner_cmd = f"cd {shlex.quote(working_dir)} && {command_str}"
            full_terminal_cmd = ['script', '-q', log_file_path, '-c', script_inner_cmd]

            if system == "Linux":
                # Assumes gnome-terminal is installed. Other terminals have different flags.
                subprocess.run(['gnome-terminal', '--wait', '--'] + full_terminal_cmd, check=True)
            else:  # macOS
                applescript_cmd = " ".join(shlex.quote(c) for c in full_terminal_cmd)
                applescript = f'tell application "Terminal" to do script "{applescript_cmd}"'
                subprocess.run(['osascript', '-e', applescript], check=True)
                # macOS needs a manual prompt to wait for the user to finish
                input("\n>> The interactive session has started in a new terminal. Please complete it, then press Enter here to continue...")
        else:
            print(f"Unsupported OS for interactive mode: {system}.")
            return False
        return True
    except (FileNotFoundError, subprocess.CalledProcessError) as e:
        print(f"\n[Error] Could not open interactive terminal: {e}.")
        print("Please ensure a standard terminal is available (e.g., PowerShell, gnome-terminal).")
        return False


# Add a function to handle GUI interactive input
def get_user_inputs_gui(file_name: str, prompts: List[str], types: List[str] | None = None) -> List[str]:
    """
    Function to be called by GUI to get user inputs for interactive programs.
    Returns a list of user inputs.
    """
    # This will be overridden by GUI if needed
    return []


def run_interpreted_code(source_file: Path, gui_input_callback=None) -> Tuple[str | None, str | None, str | None]:
    """
    Handles execution for interpreted languages like Python and JavaScript.
    """
    ext = source_file.suffix.lower()
    command = []
    if ext == '.py':
        command = ['python3', str(source_file)]
    elif ext == '.js':
        command = ['node', str(source_file)]

    # Check if the program needs interactive input
    try:
        code_content = source_file.read_text(encoding='utf-8', errors='ignore')
        if is_interactive(code_content, ext):
            print(f"\n📝 Program '{source_file.name}' requires input. Please provide values:")
            
            # For Python, look for input() calls with better pattern matching
            if ext == '.py':
                # Extract input() calls with their prompts
                input_patterns = re.findall(r'input\s*\(\s*["\']([^"\']*)["\']?\s*\)', code_content)
                user_inputs = []
                input_counter = 1
                
                # If GUI callback is provided, use it instead of console input
                if gui_input_callback:
                    prompts = []
                    for prompt in input_patterns:
                        if prompt.strip():
                            prompts.append(prompt.strip().rstrip(':').strip())
                        else:
                            prompts.append(f"Input {len(prompts) + 1}")
                    
                    if not prompts:
                        prompts = ["General Input"]
                    
                    # Call the GUI callback with the file name and list of prompts
                    user_inputs = gui_input_callback(file_name=source_file.name, prompts=prompts)
                    
                    # If user cancels input, return gracefully
                    if user_inputs is None:
                        return "User canceled input", None, None
                else:
                    # Console input (original logic)
                    for prompt in input_patterns:
                        if prompt.strip():
                            # Clean up the prompt (remove common prompt endings)
                            clean_prompt = prompt.strip().rstrip(':').strip()
                            value = input(f"{clean_prompt}: ")
                        else:
                            value = input(f"Input {input_counter}: ")
                        user_inputs.append(value)
                        input_counter += 1
                    
                    # If no specific patterns found, ask for general input
                    if not user_inputs:
                        print("Program appears to need input. Please provide values (one per line, press Enter twice when done):")
                        while True:
                            value = input(f"Input {len(user_inputs) + 1} (or press Enter to finish): ")
                            if not value:
                                break
                            user_inputs.append(value)
                
                # Run with user input
                input_data = '\n'.join(user_inputs) + '\n' if user_inputs else '\n'
                result = subprocess.run(command, input=input_data, capture_output=True, text=True, cwd=source_file.parent)
                
                # Format output to show both prompts and user responses (similar to C/C++)
                output_lines = []
                if result.stdout:
                    stdout_lines = result.stdout.strip().split('\n')
                    input_lines = user_inputs
                    input_index = 0
                    
                    # Process and replace format specifiers in the output
                    for line in stdout_lines:
                        # Check if line contains unprocessed format specifiers
                        if '%d' in line or '%s' in line or '%f' in line or '%c' in line or '%i' in line:
                            # Replace the format specifier with the actual user input if available
                            if input_index < len(input_lines):
                                try:
                                    # Try to safely format the line
                                    user_value = input_lines[input_index]
                                    # Replace common format specifiers with the user's input value
                                    processed_line = line
                                    processed_line = re.sub(r'%d|%i', user_value, processed_line)
                                    processed_line = re.sub(r'%s', user_value, processed_line)
                                    processed_line = re.sub(r'%f|%lf', user_value, processed_line)
                                    processed_line = re.sub(r'%c', user_value, processed_line)
                                    output_lines.append(processed_line)
                                except Exception:
                                    # If formatting fails, just append the original line
                                    output_lines.append(line)
                            else:
                                output_lines.append(line)
                        else:
                            output_lines.append(line)
                            
                        # If this line looks like a prompt (ends with : or contains input-related words), add the user input
                        if (line.strip().endswith(':') or 'enter' in line.lower() or 'input' in line.lower() or line.strip().endswith('?')) and input_index < len(input_lines):
                            output_lines.append(f"User entered: {input_lines[input_index]}")
                            input_index += 1
                    
                    formatted_output = '\n'.join(output_lines)
                else:
                    formatted_output = f"Program executed with user inputs: {', '.join(user_inputs)}"
                
                return formatted_output, result.stderr, None
            else:
                # For other languages, run normally
                result = subprocess.run(command, capture_output=True, text=True, cwd=source_file.parent)
                return result.stdout, result.stderr, None
        else:
            # Non-interactive execution
            result = subprocess.run(command, capture_output=True, text=True, cwd=source_file.parent)
            return result.stdout, result.stderr, None
    except Exception as e:
        return None, f"Error during execution: {e}", None


def run_c_cpp_code(source_file: Path, gui_input_callback=None) -> Tuple[str | None, str | None, str | None]:
    """
    Handles the compile-run-cleanup cycle for C and C++ files.
    With caching to avoid recompiling unchanged files.
    """
    compiler = 'g++' if source_file.suffix.lower() == '.cpp' else 'gcc'
    temp_dir = Path(tempfile.gettempdir())
    
    # Create a unique filename based on source file path and modification time
    source_stat = source_file.stat()
    source_modified_time = int(source_stat.st_mtime)
    source_size = source_stat.st_size
    source_hash = f"{source_file.stem}_{source_modified_time}_{source_size}"
    
    # Use hashed filename for caching
    exe_file = temp_dir / f"c2w_cache_{source_hash}"
    if platform.system() == "Windows":
        exe_file = exe_file.with_suffix('.exe')

    # Read the source code to check for input requirements
    code_content = source_file.read_text(encoding='utf-8', errors='ignore')
    
    # Check if we already have a compiled version
    recompile_needed = not exe_file.exists()
    
    try:
        # Only compile if the cached version doesn't exist
        if recompile_needed:
            # 1. Compile with optimizations for speed
            compile_cmd = [compiler, str(source_file), '-o', str(exe_file), '-O2']
            compile_result = subprocess.run(compile_cmd, capture_output=True, text=True)
            if compile_result.returncode != 0:
                return None, f"Compilation failed:\n{compile_result.stderr}", None

        # 2. Run - Check if it needs input
        run_cmd = [str(exe_file)]
        
        if is_interactive(code_content, source_file.suffix.lower()):
            # For interactive C/C++ programs, ask user for input values
            if not gui_input_callback:
                print(f"\n📝 Program '{source_file.name}' requires input. Please provide values:")
            
            # Extract printf statements to understand what inputs are expected
            printf_patterns = re.findall(r'printf\s*\(\s*"([^"]*)"', code_content)
            scanf_patterns = re.findall(r'scanf\s*\(\s*"([^"]*)"', code_content)
            
            user_inputs = []
            
            if gui_input_callback:
                # GUI mode - prepare prompts and types
                prompts = []
                types = []
                
                for i, scanf_format in enumerate(scanf_patterns):
                    prompt = f"Input {i+1}"
                    input_type = "text"
                    
                    # Try to find corresponding printf statement for better prompting
                    if i < len(printf_patterns):
                        prev_printf = printf_patterns[i]
                        if prev_printf.strip():
                            prompt = prev_printf.strip().rstrip(':').strip()
                    
                    # Determine input type
                    if '%d' in scanf_format or '%i' in scanf_format:
                        input_type = "integer"
                    elif '%f' in scanf_format or '%lf' in scanf_format:
                        input_type = "decimal"
                    elif '%c' in scanf_format:
                        input_type = "single character"
                    elif '%s' in scanf_format:
                        input_type = "text"
                    
                    prompts.append(f"{prompt} ({input_type})")
                    types.append(input_type)
                
                if not prompts:
                    prompts = ["General Input (text)"]
                    types = ["text"]
                
                # Call the GUI callback with the file name and list of prompts
                user_inputs = gui_input_callback(file_name=source_file.name, prompts=prompts)
                
                # If user cancels input, return gracefully
                if user_inputs is None:
                    return "User canceled input", None, None
            else:
                # Console mode (original logic)
                input_counter = 1
                
                for scanf_format in scanf_patterns:
                    prompt = f"Input {input_counter}"
                    
                    # Try to find corresponding printf statement for better prompting
                    if input_counter <= len(printf_patterns):
                        prev_printf = printf_patterns[input_counter - 1]
                        if prev_printf.strip():
                            prompt = f"{prev_printf.strip()}"
                    
                    # Determine input type and ask accordingly
                    if '%d' in scanf_format or '%i' in scanf_format:
                        value = input(f"{prompt} (integer): ")
                    elif '%f' in scanf_format or '%lf' in scanf_format:
                        value = input(f"{prompt} (decimal): ")
                    elif '%c' in scanf_format:
                        value = input(f"{prompt} (single character): ")
                    elif '%s' in scanf_format:
                        value = input(f"{prompt} (text): ")
                    else:
                        value = input(f"{prompt}: ")
                    
                    user_inputs.append(value)
                    input_counter += 1
                
                # If no specific scanf patterns found but program is interactive, ask for general input
                if not user_inputs:
                    print("Program appears to need input. Please provide values (one per line, press Enter twice when done):")
                    while True:
                        value = input(f"Input {len(user_inputs) + 1} (or press Enter to finish): ")
                        if not value:
                            break
                        user_inputs.append(value)
            
            # Combine all inputs with newlines
            input_data = '\n'.join(user_inputs) + '\n' if user_inputs else '\n'
            
            # Run with user-provided input
            run_result = subprocess.run(run_cmd, input=input_data, capture_output=True, text=True)
            
            # Format output to show both prompts and user responses
            output_lines = []
            if run_result.stdout:
                stdout_lines = run_result.stdout.strip().split('\n')
                input_lines = user_inputs
                input_index = 0
                
                # Process and replace format specifiers in the output
                for line in stdout_lines:
                    # Check if line contains unprocessed format specifiers (common in C/C++ when printing user input)
                    if '%d' in line or '%s' in line or '%f' in line or '%c' in line or '%i' in line:
                        # Replace the format specifier with the actual user input if available
                        if input_index < len(input_lines):
                            try:
                                # Try to safely format the line
                                user_value = input_lines[input_index]
                                # Replace common format specifiers with the user's input value
                                processed_line = line
                                processed_line = re.sub(r'%d|%i', user_value, processed_line)
                                processed_line = re.sub(r'%s', user_value, processed_line)
                                processed_line = re.sub(r'%f|%lf', user_value, processed_line)
                                processed_line = re.sub(r'%c', user_value, processed_line)
                                output_lines.append(processed_line)
                            except Exception:
                                # If formatting fails, just append the original line
                                output_lines.append(line)
                        else:
                            output_lines.append(line)
                    else:
                        output_lines.append(line)
                    
                    # If this line looks like a prompt (ends with : or contains "enter"), add the user input
                    if (line.strip().endswith(':') or 'enter' in line.lower() or 'input' in line.lower()) and input_index < len(input_lines):
                        output_lines.append(f"User entered: {input_lines[input_index]}")
                        input_index += 1
                
                formatted_output = '\n'.join(output_lines)
            else:
                formatted_output = f"Program executed with user inputs: {', '.join(user_inputs)}"
                
            return formatted_output, run_result.stderr, None
        else:
            # Non-interactive execution
            run_result = subprocess.run(run_cmd, capture_output=True, text=True)
            return run_result.stdout, run_result.stderr, None

    except Exception as e:
        return None, f"Error during execution: {e}", None
    finally:
        # 3. Cleanup
        if exe_file and exe_file.exists():
            exe_file.unlink()


def run_java_code(source_file: Path, gui_input_callback=None) -> Tuple[str | None, str | None, str | None]:
    """
    Handles the compile-run-cleanup cycle for Java files.
    """
    source_dir = source_file.parent
    class_file = source_dir / f"{source_file.stem}.class"

    try:
        # 1. Compile (the current working directory is the source directory)
        compile_cmd = ['javac', str(source_file)]
        compile_result = subprocess.run(compile_cmd, capture_output=True, text=True, cwd=source_dir)
        if compile_result.returncode != 0:
            return None, f"Compilation failed:\n{compile_result.stderr}", None

        # Read the source code to check for input requirements
        code_content = source_file.read_text(encoding='utf-8', errors='ignore')

        # 2. Run (using -cp to specify the classpath)
        run_cmd = ['java', '-cp', str(source_dir), source_file.stem]
        
        if is_interactive(code_content, source_file.suffix.lower()):
            # For interactive Java programs, ask user for input values
            if not gui_input_callback:
                print(f"\n📝 Program '{source_file.name}' requires input. Please provide values:")
            
            # Extract System.out.println statements to understand what inputs are expected
            println_patterns = re.findall(r'System\.out\.print(?:ln)?\s*\(\s*"([^"]*)"', code_content)
            scanner_patterns = re.findall(r'scanner\.next(?:Int|Double|Line|Float)\(\)', code_content)
            
            user_inputs = []
            
            if gui_input_callback:
                # GUI mode - prepare prompts
                prompts = []
                
                for i, scanner_call in enumerate(scanner_patterns):
                    prompt = f"Input {i+1}"
                    
                    # Try to find corresponding println statement for better prompting
                    if i < len(println_patterns):
                        prev_println = println_patterns[i]
                        if prev_println.strip():
                            prompt = prev_println.strip().rstrip(':').strip()
                    
                    # Determine input type based on scanner method
                    input_type = "text"
                    if "nextInt" in scanner_call:
                        input_type = "integer"
                    elif "nextDouble" in scanner_call or "nextFloat" in scanner_call:
                        input_type = "decimal"
                    elif "nextLine" in scanner_call:
                        input_type = "text"
                    
                    prompts.append(f"{prompt} ({input_type})")
                
                if not prompts:
                    prompts = ["General Input (text)"]
                
                # Call the GUI callback with the file name and list of prompts
                user_inputs = gui_input_callback(file_name=source_file.name, prompts=prompts)
                
                # If user cancels input, return gracefully
                if user_inputs is None:
                    return "User canceled input", None, None
            else:
                # Console mode (original logic)
                input_counter = 1
                
                for scanner_call in scanner_patterns:
                    prompt = f"Input {input_counter}"
                    
                    # Try to find corresponding println statement for better prompting
                    if input_counter <= len(println_patterns):
                        prev_println = println_patterns[input_counter - 1]
                        if prev_println.strip():
                            prompt = f"{prev_println.strip()}"
                    
                    # Determine input type based on scanner method
                    if "nextInt" in scanner_call:
                        value = input(f"{prompt} (integer): ")
                    elif "nextDouble" in scanner_call or "nextFloat" in scanner_call:
                        value = input(f"{prompt} (decimal): ")
                    else:
                        value = input(f"{prompt} (text): ")
                    
                    user_inputs.append(value)
                    input_counter += 1
                
                # If no specific scanner patterns found but program is interactive, ask for general input
                if not user_inputs:
                    print("Program appears to need input. Please provide values (one per line, press Enter twice when done):")
                    while True:
                        value = input(f"Input {len(user_inputs) + 1} (or press Enter to finish): ")
                        if not value:
                            break
                        user_inputs.append(value)
            
            # Combine all inputs with newlines
            input_data = '\n'.join(user_inputs) + '\n' if user_inputs else '\n'
            
            # Run with user-provided input
            run_result = subprocess.run(run_cmd, input=input_data, capture_output=True, text=True, cwd=source_dir)
            
            # Format output to show both prompts and user responses
            output_lines = []
            if run_result.stdout:
                stdout_lines = run_result.stdout.strip().split('\n')
                input_lines = user_inputs
                input_index = 0
                
                # Process and replace format specifiers in the output
                for line in stdout_lines:
                    # Check if line contains unprocessed format specifiers
                    if '%d' in line or '%s' in line or '%f' in line or '%c' in line or '%i' in line:
                        # Replace the format specifier with the actual user input if available
                        if input_index < len(input_lines):
                            try:
                                # Try to safely format the line
                                user_value = input_lines[input_index]
                                # Replace common format specifiers with the user's input value
                                processed_line = line
                                processed_line = re.sub(r'%d|%i', user_value, processed_line)
                                processed_line = re.sub(r'%s', user_value, processed_line)
                                processed_line = re.sub(r'%f|%lf', user_value, processed_line)
                                processed_line = re.sub(r'%c', user_value, processed_line)
                                output_lines.append(processed_line)
                            except Exception:
                                # If formatting fails, just append the original line
                                output_lines.append(line)
                        else:
                            output_lines.append(line)
                    else:
                        output_lines.append(line)
                    
                    # If this line looks like a prompt (ends with : or contains "enter"), add the user input
                    if (line.strip().endswith(':') or 'enter' in line.lower() or 'input' in line.lower()) and input_index < len(input_lines):
                        output_lines.append(f"User entered: {input_lines[input_index]}")
                        input_index += 1
                
                formatted_output = '\n'.join(output_lines)
            else:
                formatted_output = f"Program executed with user inputs: {', '.join(user_inputs)}"
                
            return formatted_output, run_result.stderr, None
        else:
            # Non-interactive execution
            run_result = subprocess.run(run_cmd, capture_output=True, text=True, cwd=source_dir)
            return run_result.stdout, run_result.stderr, None

    finally:
        # 3. Cleanup
        if class_file.exists():
            class_file.unlink()


def handle_matplotlib_plot(source_file: Path, code_content: str) -> Tuple[str | None, str | None, str | None]:
    """
    Handles Python scripts with matplotlib to save the plot instead of showing it.
    """
    temp_dir = Path(tempfile.gettempdir())
    image_path = temp_dir / "matplotlib_output.png"
    save_path_str = str(image_path).replace('\\', '/')

    # Inject code to save the figure
    injected_code = "import matplotlib\nmatplotlib.use('Agg')\n"
    code_content_mod = re.sub(r'(plt\.show\(\s*\))', f"plt.savefig('{save_path_str}')\n# \\1", code_content)
    if f"plt.savefig('{save_path_str}')" not in code_content_mod:
        code_content_mod += f"\nimport matplotlib.pyplot as plt\nplt.savefig('{save_path_str}')"

    # Run the modified script via a temporary file or command string
    result = subprocess.run(['python3', '-c', injected_code + code_content_mod], capture_output=True, text=True)
    
    if image_path.exists():
        return result.stdout, result.stderr, str(image_path)
    else:
        return result.stdout, result.stderr, None


def run_code(file_path_str: str, gui_input_callback=None) -> Tuple[str | None, str | None, str | None]:
    """
    Acts as a router to dispatch the file to the correct language-specific handler.
    Executes a given code file. Handles both interactive and non-interactive cases
    for compiled and interpreted languages.
    """
    source_file = Path(file_path_str).resolve()
    if not source_file.is_file():
        return None, f"Error: Source file not found at '{source_file}'", None

    ext = source_file.suffix.lower()
    
    try:
        code_content = source_file.read_text(encoding='utf-8', errors='ignore')
    except Exception as e:
        return None, f"Error reading file: {e}", None

    # All language-specific interactive handling is done in their respective handlers now
    # No need for special Java case here

    # The new router logic using a dictionary
    language_handlers = {
        '.py': run_interpreted_code,
        '.js': run_interpreted_code,
        '.c': run_c_cpp_code,
        '.cpp': run_c_cpp_code,
        '.java': run_java_code
    }

    handler = language_handlers.get(ext)

    if handler:
        # Check for matplotlib before calling the handler
        if ext == '.py':
            if re.search(r'import\s+(matplotlib\.pyplot|seaborn)', code_content):
                return handle_matplotlib_plot(source_file, code_content)

        # Call the appropriate handler (e.g., run_c_cpp_code)
        if ext in ['.py', '.js', '.c', '.cpp', '.java']:
            return handler(source_file, gui_input_callback)
        else:
            return handler(source_file)
    else:
        # For unsupported files like HTML, CSS
        return None, None, None


def handle_interactive_execution(source_file: Path, ext: str) -> Tuple[str | None, str | None, str | None]:
    """
    Handles interactive execution for languages that require user input.
    """
    source_dir = source_file.parent
    temp_dir = Path(tempfile.gettempdir())
    
    # Build commands based on language
    compile_cmd = None
    run_cmd = []
    cleanup_files = []
    
    if ext == '.py':
        run_cmd = ['python3', str(source_file)]
    elif ext == '.js':
        run_cmd = ['node', str(source_file)]
    elif ext in ['.c', '.cpp']:
        compiler = 'g++' if ext == '.cpp' else 'gcc'
        exe_file = temp_dir / f"temp_program_{source_file.stem}"
        if platform.system() == "Windows":
            exe_file = exe_file.with_suffix('.exe')
        
        compile_cmd = [compiler, str(source_file), '-o', str(exe_file)]
        run_cmd = [str(exe_file)]
        cleanup_files.append(exe_file)
    elif ext == '.java':
        compile_cmd = ['javac', str(source_file)]
        run_cmd = ['java', '-cp', str(source_dir), source_file.stem]
        cleanup_files.append(source_dir / f"{source_file.stem}.class")
    
    try:
        # Compile if necessary
        if compile_cmd:
            compile_result = subprocess.run(compile_cmd, capture_output=True, text=True)
            if compile_result.returncode != 0:
                return None, f"Compilation failed:\n{compile_result.stderr}", None

        # Handle interactive execution
        if not run_cmd:
            return None, "No execution command defined for this file type", None
            
        fd, log_file_path = tempfile.mkstemp(text=True, suffix=".log")
        os.close(fd)
        
        working_dir = str(source_dir if ext == '.java' else temp_dir)
        success = open_interactive_terminal(run_cmd, working_dir, log_file_path)

        if success:
            log_content = Path(log_file_path).read_text(encoding='utf-8', errors='ignore')
            Path(log_file_path).unlink()
            # Basic cleaning of terminal control characters
            log_content = re.sub(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])', '', log_content)
            return log_content.strip(), None, None
        else:
            if Path(log_file_path).exists(): 
                Path(log_file_path).unlink()
            return None, "Failed to open or log the interactive session.", None
    
    finally:
        # Cleanup
        for f in cleanup_files:
            if f and Path(f).exists():
                Path(f).unlink()


# --- UNCHANGED FUNCTIONS ---

def generate_full_page_screenshots(file_path: str, viewport_height: int = 1394) -> List[str]:
    # Lazy import - only import selenium when needed
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service
    from webdriver_manager.chrome import ChromeDriverManager
    
    options = webdriver.ChromeOptions()
    if BROWSER_BINARY_LOCATION:
        options.binary_location = BROWSER_BINARY_LOCATION

    # Performance optimized options
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-gpu')
    options.add_argument('--disable-extensions')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--disable-setuid-sandbox')
    options.add_argument('--disable-features=TranslateUI')
    options.add_argument('--disable-web-security')
    options.add_argument('--dns-prefetch-disable')
    options.add_argument('--disable-infobars')
    options.add_argument('--blink-settings=imagesEnabled=false') # Disable images for speed

    # Create a single driver instance and reuse it
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

    abs_path = os.path.abspath(file_path)
    file_url = 'file:///' + abs_path.replace('\\', '/')

    driver.set_window_size(1920, viewport_height)
    driver.get(file_url)
    time.sleep(1)  # Reduced wait time

    screenshot_paths = []
    page_height = driver.execute_script("return document.body.scrollHeight")
    current_scroll_position = 0
    screenshot_counter = 0

    while current_scroll_position < page_height:
        temp_dir = tempfile.gettempdir()
        screenshot_name = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(file_path))[0]}_screenshot_{screenshot_counter:02d}.png")
        
        driver.save_screenshot(screenshot_name)
        screenshot_paths.append(screenshot_name)
        screenshot_counter += 1

        current_scroll_position += viewport_height
        driver.execute_script(f"window.scrollTo(0, {current_scroll_position});")
        time.sleep(0.5)

        page_height = driver.execute_script("return document.body.scrollHeight")
        
        if current_scroll_position >= page_height and screenshot_counter > 1:
            break
        elif current_scroll_position + viewport_height > page_height and current_scroll_position < page_height:
            driver.execute_script(f"window.scrollTo(0, {page_height});")
            time.sleep(0.5)
            screenshot_name = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(file_path))[0]}_screenshot_{screenshot_counter:02d}.png")
            driver.save_screenshot(screenshot_name)
            screenshot_paths.append(screenshot_name)
            break

    driver.quit()
    return screenshot_paths


def add_code_block(doc: Any, code_text: str, is_error: bool = False):
    """
    Adds a formatted code block to the Word document with proper styling.
    """
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    color = "FADBD8" if is_error else "D9D9D9" # Light red for errors, gray for normal
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

    paragraph = cell.paragraphs[0]
    
    # Ensure proper formatting with newlines preserved
    formatted_text = code_text.strip() if code_text else "[No output]"
    
    run = paragraph.add_run(formatted_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # Add some padding around the code block
    paragraph.space_before = Pt(6)
    paragraph.space_after = Pt(6)


def set_fixed_cell_width(cell, width_in_inches: float):
    cell.width = Inches(width_in_inches)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(int(width_in_inches * 1440)))
    tcPr.append(tcW)


def create_student_info_table(doc: Any, metadata: Dict[str, str]):
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    for row in table.rows:
        set_fixed_cell_width(row.cells[0], 2)
        set_fixed_cell_width(row.cells[1], 2.3)

    info = {
        'Name': metadata.get('name', 'N/A'),
        'Registration Number': metadata.get('reg_no', 'N/A'),
        'Semester': metadata.get('semester', 'N/A'),
        'Group': metadata.get('group', 'N/A')
    }

    for i, (key, value) in enumerate(info.items()):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_page_border(section):
    sectPr = section._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    border_attributes = {
        'w:val': 'single', 'w:sz': '24',
        'w:space': '24', 'w:color': 'auto'
    }
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        for attr, value in border_attributes.items():
            border_el.set(qn(attr), value)
        pg_borders.append(border_el)
    sectPr.append(pg_borders)


def create_word_doc(files: List[str], metadata: Dict[str, str], output_dir: str = ".", gui_input_callback=None):
    """
    Creates a formatted Word document with code, output, and metadata.
    Optimized for better performance.
    """
    import concurrent.futures
    
    os.makedirs(output_dir, exist_ok=True)
    
    # Convert all paths to absolute paths
    files = [os.path.abspath(f) for f in files]
    output_dir = os.path.abspath(output_dir)

    # Create document with minimized settings
    doc = Document()
    section = doc.sections[0]
    add_page_border(section)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Main heading with better styling
    heading = doc.add_heading(f"{metadata['subject']} - Experiment {metadata['experiment_no']}", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.font.size = Pt(28)
    run.font.bold = True
    
    # Add a subtitle with student info summary
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run(f"Submitted by: {metadata.get('name', 'N/A')} | Registration: {metadata.get('reg_no', 'N/A')}")
    subtitle_run.font.size = Pt(12)
    subtitle_run.italic = True
    
    # Add some spacing after header
    doc.add_paragraph()

    for i, file_path in enumerate(files):
        if not os.path.isfile(file_path):
            print(f"Skipping: {file_path} (file not found)")
            continue

        ext = os.path.splitext(file_path)[1].lower()

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                code_text = f.read()
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            continue

        if i > 0:
            # Add a horizontal line as a divider between files instead of a page break
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)  # Add some space before the line
            p.paragraph_format.space_after = Pt(20)   # Add some space after the line
            p_fmt = p.paragraph_format
            p_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p_run = p.add_run()
            p_run.add_break(WD_BREAK.LINE)
            
            # Add a horizontal line
            p.runs[0].add_tab()
            p.runs[0].add_text('_' * 50)  # 50 underscore characters for the line

        # File header with better formatting
        file_heading = doc.add_heading(f'File: {os.path.basename(file_path)}', level=2)
        file_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Source code section
        doc.add_heading('Source Code:', level=3)
        add_code_block(doc, code_text)
        
        # Add spacing between sections
        doc.add_paragraph()

        if ext in ['.html', '.css']: 
            try:
                if ext == '.html':
                    screenshot_paths = generate_full_page_screenshots(file_path)
                    doc.add_heading('Frontend Preview:', level=3)
                    if screenshot_paths:
                        for s_path in screenshot_paths:
                            # Add the screenshot with proper sizing
                            try:
                                doc.add_picture(s_path, width=Inches(6.5))
                                doc.add_paragraph()  # Add spacing after image
                                os.remove(s_path)
                            except Exception as e:
                                print(f"Error adding screenshot: {e}")
                    else:
                        doc.add_paragraph("[No screenshots generated]")
                else: # For .css
                    doc.add_heading('CSS Information:', level=3)
                    doc.add_paragraph("CSS files provide styling for web pages and have no direct executable output.")
            except Exception as e:
                doc.add_heading('Frontend Error:', level=3)
                add_code_block(doc, str(e), is_error=True)
        else:
            try:
                output_text, error_text, image_path = run_code(file_path, gui_input_callback)
                
                # Output section with better formatting
                doc.add_heading('Program Output:', level=3)
                if output_text and output_text.strip():
                    add_code_block(doc, output_text)
                else:
                    doc.add_paragraph("[No standard output produced]")
                
                # Add spacing
                doc.add_paragraph()

                # Image output section
                if image_path and os.path.exists(image_path):
                    doc.add_heading('Generated Plot/Image:', level=3)
                    try:
                        doc.add_picture(image_path, width=Inches(6.5))
                        doc.add_paragraph()  # Add spacing after image
                        os.remove(image_path)
                    except Exception as e:
                        print(f"Error adding image: {e}")
                
                # Error section (if any)
                if error_text and error_text.strip():
                    doc.add_heading('Errors/Warnings:', level=3)
                    add_code_block(doc, error_text, is_error=True)
                    doc.add_paragraph()  # Add spacing after errors
                    
            except Exception as e:
                doc.add_heading('Execution Error:', level=3)
                add_code_block(doc, str(e), is_error=True)

    # Add final page break before submission info
    doc.add_page_break()
    
    # Submission information section
    doc.add_heading('Submission Information', level=2)
    create_student_info_table(doc, metadata)

    # Generate filename and save
    name_prefix = metadata.get('name', 'Student').split()[0].replace(' ', '_')
    subject_clean = metadata.get('subject', 'Subject').replace(' ', '_')
    exp_no = metadata.get('experiment_no', '1')
    output_filename = f"{name_prefix}_{subject_clean}_Exp{exp_no}.docx"
    output_docx = os.path.join(output_dir, output_filename)
    
    # Optimized saving
    try:
        doc.save(output_docx)
        print(f"\n✅ Final report saved as {output_docx}")
        
        # Clean up any temp files if they exist
        import glob
        temp_files = glob.glob(os.path.join(tempfile.gettempdir(), "c2w_*"))
        for temp_file in temp_files:
            try:
                if os.path.isfile(temp_file) and os.path.getmtime(temp_file) < time.time() - 7200:  # Older than 2 hours
                    os.unlink(temp_file)
            except:
                pass
                
        return output_docx
    except Exception as e:
        print(f"Error saving document: {e}")
        # Try to save in a different location as backup
        backup_path = os.path.join(tempfile.gettempdir(), output_filename)
        try:
            doc.save(backup_path)
            print(f"Backup saved to {backup_path}")
            return backup_path
        except:
            print("Could not save backup document")
            return None


def main():

    print("-" * 40)
    print("📝 Enter Assignment Info:")
    metadata = {
        'name': input("Enter your name: "),
        'subject': input("Enter subject name: "),
        'reg_no': input("Enter registration number: "),
        'group': input("Enter group: "),
        'semester': input("Enter semester: "),
        'experiment_no': input("Enter experiment number: ")
    }

    if len(sys.argv) < 2:
        print("\nUsage: python code_to_word_report.py <file1> <file2> ...")
        return

    files = [os.path.abspath(f) for f in sys.argv[1:]]
    create_word_doc(files, metadata, ".", None)


if __name__ == '__main__':
    main()